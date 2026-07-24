from __future__ import annotations
import argparse,csv,html,json,math
from pathlib import Path
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font,PatternFill,Alignment
from openpyxl.utils import get_column_letter
from validate_process_spec import validate
from generate_presentation import generate_presentation

COLORS=['D9EAF7','E2F0D9','FCE4D6','E4DFEC','FFF2CC','DDEBF7']
def autofit(ws):
    for i in range(1,ws.max_column+1):
        vals=[len(str(ws.cell(r,i).value or '')) for r in range(1,min(ws.max_row,100)+1)];ws.column_dimensions[get_column_letter(i)].width=min(45,max(12,max(vals or [10])+2))
def header(ws,row=1):
    for c in ws[row]:c.fill=PatternFill('solid',fgColor='1F4E78');c.font=Font(color='FFFFFF',bold=True);c.alignment=Alignment(wrap_text=True)
def add_table(doc,headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid'
    for i,h in enumerate(headers):t.rows[0].cells[i].text=str(h)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):cells[i].text='' if v is None else str(v)
    return t

def generate_svg(spec,path):
    steps=spec['steps'];byid={s['id']:s for s in steps};owners=[]
    for s in steps:
        if s.get('owner') not in owners:owners.append(s.get('owner'))
    lane_w=260;box_w=200;box_h=72;gap_y=55;margin=40;top=95;width=margin*2+lane_w*len(owners);height=top+len(steps)*(box_h+gap_y)+80
    lane_x={o:margin+i*lane_w for i,o in enumerate(owners)};pos={};order={s['id']:i for i,s in enumerate(steps)}
    parts=[f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">', '<defs><marker id="arrow" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#5B6573"/></marker></defs>', '<rect width="100%" height="100%" fill="white"/>',f'<text x="{margin}" y="35" font-family="Arial" font-size="24" font-weight="bold">{html.escape(spec["process_name"])}</text>']
    for i,o in enumerate(owners):
        x=lane_x[o];parts.append(f'<rect x="{x}" y="55" width="{lane_w}" height="{height-75}" fill="#{COLORS[i%len(COLORS)]}" fill-opacity="0.28" stroke="#B8C2CC"/>');parts.append(f'<text x="{x+12}" y="82" font-family="Arial" font-size="15" font-weight="bold">{html.escape(str(o))}</text>')
    for idx,s in enumerate(steps):
        x=lane_x[s.get('owner')]+30;y=top+idx*(box_h+gap_y);pos[s['id']]=(x,y)
        typ=s.get('type','task')
        if typ=='decision':
            cx=x+box_w/2;cy=y+box_h/2;pts=f'{cx},{y} {x+box_w},{cy} {cx},{y+box_h} {x},{cy}';parts.append(f'<polygon points="{pts}" fill="#FFF2CC" stroke="#7F6000" stroke-width="2"/>')
        elif typ in ('start','end'):
            cx=x+box_w/2;cy=y+box_h/2;parts.append(f'<ellipse cx="{cx}" cy="{cy}" rx="{box_w/2}" ry="{box_h/2}" fill="#E2EFDA" stroke="#375623" stroke-width="2"/>')
        else:parts.append(f'<rect x="{x}" y="{y}" rx="10" ry="10" width="{box_w}" height="{box_h}" fill="#FFFFFF" stroke="#1F4E78" stroke-width="2"/>')
        label=f"{s['id']}  {s['name']}";words=label.split();lines=[];cur=''
        for w in words:
            if len(cur)+len(w)+1>28:lines.append(cur);cur=w
            else:cur=(cur+' '+w).strip()
        if cur:lines.append(cur)
        for li,line in enumerate(lines[:3]):parts.append(f'<text x="{x+box_w/2}" y="{y+28+li*17}" text-anchor="middle" font-family="Arial" font-size="12" font-weight="bold">{html.escape(line)}</text>')
    for s in steps:
        x,y=pos[s['id']]
        targets=[]
        if s.get('type')=='decision':targets=[('Yes',s.get('yes_next')),('No',s.get('no_next'))]
        elif s.get('next'):targets=[('',s.get('next'))]
        for label,t in targets:
            if not t or t not in pos:continue
            tx,ty=pos[t];x1=x+box_w/2;y1=y+box_h;x2=tx+box_w/2;y2=ty
            dash=' stroke-dasharray="7 5"' if order.get(t,10**9)<=order.get(s['id'],-1) else ''
            parts.append(f'<path d="M {x1} {y1} C {x1} {y1+25}, {x2} {y2-25}, {x2} {y2}" fill="none" stroke="#5B6573" stroke-width="2"{dash} marker-end="url(#arrow)"/>')
            if label:parts.append(f'<text x="{(x1+x2)/2+6}" y="{(y1+y2)/2}" font-family="Arial" font-size="11">{label}</text>')
    parts.append('</svg>');path.write_text('\n'.join(parts),encoding='utf-8')

def generate_docx(spec,path):
    d=Document();sec=d.sections[0];sec.top_margin=Inches(.65);sec.bottom_margin=Inches(.65)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(spec['process_name']);r.bold=True;r.font.size=Pt(22)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run(f"Standard Operating Procedure | Version {spec.get('version','1.0')} | {spec.get('state','current').title()} State")
    for heading,text in [('Purpose',spec.get('purpose')),('Trigger',spec.get('trigger'))]:d.add_heading(heading,level=1);d.add_paragraph(text or 'TBD')
    d.add_heading('Scope',level=1);sc=spec.get('scope') or {};d.add_paragraph(f"Starts: {sc.get('starts','TBD')}");d.add_paragraph(f"Ends: {sc.get('ends','TBD')}");d.add_paragraph('Exclusions: '+(', '.join(sc.get('exclusions') or []) or 'None stated'))
    d.add_heading('Inputs and Outputs',level=1);add_table(d,['Inputs','Outputs'],[[', '.join(spec.get('inputs') or []),', '.join(spec.get('outputs') or [])]])
    d.add_heading('Roles and Responsibilities',level=1);add_table(d,['Role','Responsibility'],[[r.get('name'),r.get('description')] for r in (spec.get('roles') or [])])
    d.add_heading('Procedure',level=1);rows=[]
    for s in (spec.get('steps') or []):
        nxt=(f"Yes → {s.get('yes_next')}; No → {s.get('no_next')}" if s.get('type')=='decision' else s.get('next') or 'End')
        rows.append([s.get('id'),s.get('name'),s.get('owner'),s.get('description'),s.get('system','') or 'TBD',s.get('sla','') or 'TBD',nxt,', '.join(s.get('control_ids') or []) or 'None'])
    add_table(d,['ID','Step','Owner','Instructions','System','SLA','Next','Controls'],rows)
    d.add_heading('Controls and Evidence',level=1);add_table(d,['ID','Control','Objective','Owner','Frequency','Evidence','Type'],[[c.get(k) or 'TBD' for k in ['id','name','objective','owner','frequency','evidence','type']] for c in (spec.get('controls') or [])])
    d.add_heading('Risks',level=1);add_table(d,['ID','Risk','Impact','Mitigation'],[[r.get('id'),r.get('description'),r.get('impact'),r.get('mitigation')] for r in (spec.get('risks') or [])])
    d.add_heading('Metrics',level=1);add_table(d,['Metric','Definition','Target','Owner'],[[m.get('name'),m.get('definition'),m.get('target'),m.get('owner')] for m in (spec.get('metrics') or [])])
    d.add_heading('Assumptions and Open Questions',level=1)
    assumptions=spec.get('assumptions') or [];open_questions=spec.get('open_questions') or []
    for x in assumptions:d.add_paragraph('Assumption: '+x,style='List Bullet')
    for x in open_questions:d.add_paragraph('Open question: '+x,style='List Bullet')
    if not assumptions and not open_questions:d.add_paragraph('TBD — no assumptions or open questions captured.',style='List Bullet')
    d.save(path)

def generate_xlsx(spec,path):
    wb=Workbook();r=wb.active;r.title='RACI';roles=[x.get('name') for x in (spec.get('roles') or [])];r.append(['Step ID','Process Step']+roles)
    for s in (spec.get('steps') or []):
        row=[s.get('id'),s.get('name')]
        for role in roles:row.append('R/A' if role==s.get('owner') else 'C' if role in (s.get('consulted_roles') or []) else 'I' if role in (s.get('informed_roles') or []) else '')
        r.append(row)
    header(r);autofit(r);r.freeze_panes='C2'
    c=wb.create_sheet('Control Register');c.append(['ID','Control','Objective','Owner','Frequency','Evidence','Type','Related Steps'])
    for ctl in (spec.get('controls') or []):
        related=[s.get('id') for s in (spec.get('steps') or []) if ctl.get('id') in (s.get('control_ids') or [])]
        c.append([ctl.get(k) or 'TBD' for k in ['id','name','objective','owner','frequency','evidence','type']]+[', '.join(related) or 'TBD'])
    header(c);autofit(c)
    m=wb.create_sheet('Metrics');m.append(['Metric','Definition','Target','Owner'])
    for x in (spec.get('metrics') or []):m.append([x.get('name'),x.get('definition'),x.get('target'),x.get('owner')])
    header(m);autofit(m);wb.save(path)

def main():
    p=argparse.ArgumentParser();p.add_argument('--input',required=True,type=Path);p.add_argument('--output-dir',required=True,type=Path);a=p.parse_args();spec=json.loads(a.input.read_text(encoding='utf-8'));errors,warnings=validate(spec)
    if errors:raise SystemExit('Validation errors: '+'; '.join(errors))
    a.output_dir.mkdir(parents=True,exist_ok=True);slug=''.join(ch.lower() if ch.isalnum() else '_' for ch in spec['process_name']).strip('_')
    generate_docx(spec,a.output_dir/f'{slug}_sop.docx');generate_svg(spec,a.output_dir/f'{slug}_process_map.svg');generate_xlsx(spec,a.output_dir/f'{slug}_raci_and_controls.xlsx')
    theme_path=Path(__file__).resolve().parents[1]/'assets'/'presentation_theme.json'
    generate_presentation(spec,a.output_dir/f'{slug}_executive_process_briefing.pptx',theme_path if theme_path.exists() else None)
    with (a.output_dir/f'{slug}_improvement_backlog.csv').open('w',encoding='utf-8-sig',newline='') as f:
        fields=['title','problem','recommendation','value','effort','risk','priority','owner'];w=csv.DictWriter(f,fieldnames=fields);w.writeheader();w.writerows([{k:x.get(k,'') for k in fields} for x in (spec.get('improvements') or [])])
    lines=[f"# {spec['process_name']} — Process Pack",'',f"**Purpose:** {spec.get('purpose','')}",f"**State:** {spec.get('state','current')}",f"**Steps:** {len(spec.get('steps') or [])}",f"**Controls:** {len(spec.get('controls') or [])}",f"**Improvement opportunities:** {len(spec.get('improvements') or [])}",'**Presentation:** Executive process briefing generated as editable PowerPoint.','','## Open questions']+[f'- {x}' for x in (spec.get('open_questions') or [])]+['','## Validation warnings']+[f'- {x}' for x in warnings]
    (a.output_dir/f'{slug}_summary.md').write_text('\n'.join(lines),encoding='utf-8');print(f'Process pack written: {a.output_dir}')
if __name__=='__main__':main()
