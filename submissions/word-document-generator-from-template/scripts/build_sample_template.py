#!/usr/bin/env python3
"""Regenerate the bundled sample DOCX used by docs and tests."""

from __future__ import annotations

import argparse
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw


def _split_token(paragraph, token: str, pieces: tuple[str, ...] | None = None):
    """Add a token as several Word runs to test run-safe replacement."""
    parts = pieces or (token[:2], token[2:-2], token[-2:])
    for index, part in enumerate(parts):
        run = paragraph.add_run(part)
        if index == 1:
            run.bold = True


def _field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, result, end))


def _style_document(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for name, size, color in (
        ("Title", 26, RGBColor(31, 78, 121)),
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 12, RGBColor(68, 68, 68)),
    ):
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color


def _make_logo(path: Path) -> None:
    image = Image.new("RGB", (600, 140), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((5, 5, 595, 135), radius=24, fill=(31, 78, 121))
    draw.text((35, 45), "CONTOSO  |  DOCUMENT TEMPLATE", fill="white")
    image.save(path, "PNG")


def build_sample(output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _style_document(document)

    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    with tempfile.TemporaryDirectory() as temp_dir:
        logo = Path(temp_dir) / "logo.png"
        _make_logo(logo)
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_p.add_run().add_picture(str(logo), width=Inches(4.5))
        header_p.add_run("\nInternal | ")
        _split_token(header_p, "{{document.title}}")

        title = document.add_paragraph(style="Title")
        _split_token(title, "{{document.title}}", ("{{document.", "title", "}}"))
        subtitle = document.add_paragraph()
        subtitle.add_run("Generated deterministically from approved sources").italic = True

        document.add_heading("Document control", level=1)
        control = document.add_table(rows=5, cols=2)
        control.style = "Table Grid"
        metadata = [
            ("Type", "{{document.type}}"),
            ("Owner", "{{document.owner}}"),
            ("Version", "{{document.version}}"),
            ("Status", "{{document.status}}"),
            ("Audience", "{{document.audience}}"),
        ]
        for row, (label, token) in zip(control.rows, metadata):
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = ""
            _split_token(row.cells[1].paragraphs[0], token)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for heading, token in (
            ("Executive summary", "{{sections.executive_summary}}"),
            ("Purpose", "{{sections.purpose}}"),
            ("Scope", "{{sections.scope}}"),
        ):
            document.add_heading(heading, level=1)
            paragraph = document.add_paragraph()
            _split_token(paragraph, token)

        document.add_heading("Findings", level=1)
        findings = document.add_table(rows=2, cols=3)
        findings.style = "Table Grid"
        for cell, label in zip(
            findings.rows[0].cells, ("Finding", "Impact", "Owner")
        ):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        row = findings.rows[1]
        for cell, token in zip(
            row.cells,
            (
                "{{findings[].finding}}",
                "{{findings[].impact}}",
                "{{findings[].owner}}",
            ),
        ):
            cell.text = ""
            _split_token(cell.paragraphs[0], token)

        document.add_heading("Recommendations", level=1)
        recommendations = document.add_paragraph()
        _split_token(recommendations, "{{sections.recommendations}}")

        # A second section demonstrates that section properties and linked
        # header/footer relationships survive the fill.
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading("Appendix", level=1)
        appendix = document.add_paragraph()
        _split_token(appendix, "{{sections.appendix}}")

        for sec in document.sections:
            footer = sec.footer
            footer.is_linked_to_previous = True
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _split_token(footer_p, "{{document.version}}")
        footer_p.add_run(" | Page ")
        _field(footer_p, "PAGE")
        footer_p.add_run(" of ")
        _field(footer_p, "NUMPAGES")
        footer_p.add_run(" | ")
        _split_token(footer_p, "{{document.status}}")

        document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "output",
        nargs="?",
        default=str(
            Path(__file__).resolve().parents[1] / "assets" / "sample-template.docx"
        ),
    )
    args = parser.parse_args()
    output = Path(args.output).resolve()
    build_sample(output)
    print(output)


if __name__ == "__main__":
    main()
