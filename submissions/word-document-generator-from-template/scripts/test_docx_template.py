#!/usr/bin/env python3
"""Regression tests for deterministic DOCX template filling."""

from __future__ import annotations

import copy
import hashlib
import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

from build_sample_template import build_sample
from docx_template import (
    TemplateError,
    fill_template,
    inspect_template,
    validate_docx,
)


HERE = Path(__file__).resolve().parent
ASSETS = HERE.parent / "assets"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def _read_zip(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path, "r") as archive:
        return {name: archive.read(name) for name in archive.namelist()}


def _write_zip(path: Path, parts: dict[str, bytes]) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, value in parts.items():
            archive.writestr(name, value)


def _visible_text(xml: bytes) -> str:
    root = etree.fromstring(xml)
    return "".join(root.xpath(".//w:t/text()", namespaces=NS))


class DocxTemplateTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.template = self.root / "template.docx"
        build_sample(self.template)
        self.data = json.loads(
            (ASSETS / "sample-data.json").read_text(encoding="utf-8")
        )

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_inspect_finds_split_tokens_arrays_and_fields(self) -> None:
        manifest = inspect_template(self.template)
        self.assertIn("document.title", manifest["scalar_placeholders"])
        self.assertIn("sections.executive_summary", manifest["scalar_placeholders"])
        self.assertEqual(
            manifest["repeating_arrays"],
            {"findings": ["finding", "impact", "owner"]},
        )
        field_json = json.dumps(manifest["word_fields"])
        self.assertIn("PAGE", field_json)
        self.assertIn("NUMPAGES", field_json)
        self.assertIn("word/header1.xml", manifest["parts"])
        self.assertIn("word/footer1.xml", manifest["parts"])

    def test_fill_preserves_fields_package_parts_and_original(self) -> None:
        original_hash = hashlib.sha256(self.template.read_bytes()).hexdigest()
        original_parts = _read_zip(self.template)
        output = self.root / "nested" / "filled.docx"

        report = fill_template(self.template, self.data, output)

        self.assertTrue(output.is_file())
        self.assertEqual(
            hashlib.sha256(self.template.read_bytes()).hexdigest(), original_hash
        )
        self.assertTrue(report["field_signature_preserved"])
        self.assertEqual(report["repeated_rows"], {"findings": 3})
        self.assertEqual(report["defaulted_fields"], [])
        self.assertTrue(report["validation"]["valid_docx"])

        result_parts = _read_zip(output)
        for name, content in original_parts.items():
            if name not in {
                "word/document.xml",
                "word/header1.xml",
                "word/footer1.xml",
            }:
                self.assertEqual(
                    result_parts[name],
                    content,
                    f"Unmodified package part changed: {name}",
                )

        text = " ".join(
            _visible_text(result_parts[name])
            for name in ("word/document.xml", "word/header1.xml", "word/footer1.xml")
        )
        self.assertIn("Quarterly Operations Report", text)
        self.assertIn("Service response targets were met.", text)
        self.assertNotIn("{{", text)

        document = Document(output)
        findings_table = next(
            table for table in document.tables if table.rows[0].cells[0].text == "Finding"
        )
        self.assertEqual(len(findings_table.rows), 4)  # header + 3 items

    def test_newlines_become_word_line_breaks(self) -> None:
        output = self.root / "line-breaks.docx"
        fill_template(self.template, self.data, output)
        root = etree.fromstring(_read_zip(output)["word/document.xml"])
        self.assertGreaterEqual(len(root.xpath(".//w:br", namespaces=NS)), 2)

    def test_missing_scalar_uses_fallback_and_reports_it(self) -> None:
        data = copy.deepcopy(self.data)
        del data["document"]["audience"]
        output = self.root / "missing.docx"
        report = fill_template(self.template, data, output)
        self.assertIn("document.audience", report["defaulted_fields"])
        text = _visible_text(_read_zip(output)["word/document.xml"])
        self.assertIn("Not specified in approved sources", text)

    def test_empty_array_removes_sample_row(self) -> None:
        data = copy.deepcopy(self.data)
        data["findings"] = []
        output = self.root / "empty.docx"
        report = fill_template(self.template, data, output)
        self.assertEqual(report["repeated_rows"], {"findings": 0})
        document = Document(output)
        findings_table = next(
            table for table in document.tables if table.rows[0].cells[0].text == "Finding"
        )
        self.assertEqual(len(findings_table.rows), 1)

    def test_wrong_array_type_fails_without_output(self) -> None:
        data = copy.deepcopy(self.data)
        data["findings"] = {"finding": "not an array"}
        output = self.root / "bad-array.docx"
        with self.assertRaisesRegex(TemplateError, "requires a JSON array"):
            fill_template(self.template, data, output)
        self.assertFalse(output.exists())

    def test_complex_scalar_fails_without_output(self) -> None:
        data = copy.deepcopy(self.data)
        data["document"]["title"] = {"nested": "not supported"}
        output = self.root / "bad-scalar.docx"
        with self.assertRaisesRegex(TemplateError, "requires a scalar"):
            fill_template(self.template, data, output)
        self.assertFalse(output.exists())

    def test_input_output_collision_is_rejected(self) -> None:
        with self.assertRaisesRegex(TemplateError, "must differ"):
            fill_template(self.template, self.data, self.template)

    def test_malformed_docx_is_rejected(self) -> None:
        bad = self.root / "bad.docx"
        bad.write_text("not a zip", encoding="utf-8")
        with self.assertRaisesRegex(TemplateError, "Cannot read DOCX"):
            inspect_template(bad)

    def test_malformed_placeholder_is_rejected(self) -> None:
        parts = _read_zip(self.template)
        root = etree.fromstring(parts["word/document.xml"])
        target = next(
            node
            for node in root.xpath(".//w:t", namespaces=NS)
            if "sections.executive_summary" in (node.text or "")
        )
        target.text = "sections executive_summary"
        parts["word/document.xml"] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8"
        )
        malformed = self.root / "malformed-token.docx"
        _write_zip(malformed, parts)
        with self.assertRaisesRegex(TemplateError, "Malformed placeholder"):
            inspect_template(malformed)

    def test_unmatched_placeholder_braces_are_rejected(self) -> None:
        parts = _read_zip(self.template)
        root = etree.fromstring(parts["word/document.xml"])
        target = next(
            node
            for node in root.xpath(".//w:t", namespaces=NS)
            if "sections.executive_summary" in (node.text or "")
        )
        target.text = "sections.executive_summary"
        closing = target.getparent().getnext()
        if closing is not None:
            closing_text = closing.find(f"{{{W}}}t")
            if closing_text is not None:
                closing_text.text = ""
        parts["word/document.xml"] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8"
        )
        malformed = self.root / "unmatched-token.docx"
        _write_zip(malformed, parts)
        with self.assertRaisesRegex(TemplateError, "Malformed placeholder"):
            inspect_template(malformed)

    def test_validate_detects_removed_live_field(self) -> None:
        output = self.root / "filled.docx"
        fill_template(self.template, self.data, output)
        parts = _read_zip(output)
        footer = etree.fromstring(parts["word/footer1.xml"])
        instr = footer.xpath(".//w:instrText", namespaces=NS)[0]
        instr.getparent().remove(instr)
        parts["word/footer1.xml"] = etree.tostring(
            footer, xml_declaration=True, encoding="UTF-8"
        )
        damaged = self.root / "damaged.docx"
        _write_zip(damaged, parts)
        with self.assertRaisesRegex(TemplateError, "field signature"):
            validate_docx(damaged, template_path=self.template)

    def test_cli_inspect_fill_validate(self) -> None:
        data_path = self.root / "data.json"
        data_path.write_text(json.dumps(self.data), encoding="utf-8")
        manifest = self.root / "manifest.json"
        output = self.root / "cli-output.docx"
        summary = self.root / "summary.json"
        validation = self.root / "validation.json"
        script = HERE / "docx_template.py"

        commands = [
            [
                sys.executable,
                str(script),
                "inspect",
                str(self.template),
                "--output",
                str(manifest),
            ],
            [
                sys.executable,
                str(script),
                "fill",
                str(self.template),
                str(data_path),
                str(output),
                "--summary",
                str(summary),
            ],
            [
                sys.executable,
                str(script),
                "validate",
                str(output),
                "--template",
                str(self.template),
                "--output",
                str(validation),
            ],
        ]
        for command in commands:
            result = subprocess.run(
                command, capture_output=True, text=True, encoding="utf-8"
            )
            self.assertEqual(result.returncode, 0, result.stderr)
        for path in (manifest, output, summary, validation):
            self.assertTrue(path.exists(), path)
        self.assertTrue(
            json.loads(validation.read_text(encoding="utf-8"))["valid_docx"]
        )


if __name__ == "__main__":
    unittest.main(verbosity=2)
